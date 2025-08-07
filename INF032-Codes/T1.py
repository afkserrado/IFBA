######################################################################
# Bibliotecas
from pathlib import Path # Biblioteca para manipulação de caminhos
from docx import Document # Biblioteca para manipulação de documentos no formato .docx
import logging

# Configuração básica do logging
logging.basicConfig(
    # Nome do arquivo de log
    filename = 'erros_resumos.log',  
    
    # 'a' para acrescentar ao final da linha
    filemode = 'a', 
    
    # Data e hora - Nível do log - Mensagem do log
    format = '%(asctime)s - %(levelname)s - %(message)s', 
    
    # Define o nível mínimo de severidade do erro para ser registrado no arquivo
    level = logging.ERROR # Registra apenas erros do tipo ERROR e CRITICAL
)

######################################################################
# Funções

# Guarda o caminho das subpastas de uma pasta em uma lista
def pegarSubpastas(cmPasta):
    cmSubpastas = []

    # Tenta acessar o caminho da pasta principal
    try:
        # Verifica se a pasta existe e se corresponde a um diretório
        if cmPasta.exists() and cmPasta.is_dir():
            # Guarda o caminho de todos os arquivos e subpastas
            itens = cmPasta.iterdir()
            
            # Percorre todos os caminhos dentro da pasta principal
            for item in itens:
                # Tenta acessar o caminho de uma subpasta ou arquivo
                try: 
                    # Verifica se o caminho corresponde a uma pasta
                    if item.is_dir(): 
                        nome = item.name[:8] 
                        partes = nome.split()

                        # Verifica se os oito primeiros caracteres do nome da subpasta correspondem ao padrão AA MM DD
                        if len(partes) == 3 and all(parte.isdigit() for parte in partes):
                            cmSubpastas.append(item)

                # Em caso de qualquer erro, registra no log
                except Exception as erro:
                    logging.error(f"Erro ao acessar a subpasta {item}: {erro}")
                    continue
 
        else:
            print("O caminho não existe ou não é uma pasta.")
        
    # Em caso de qualquer erro, registra no log
    except Exception as erro:
        logging.error(f"Erro ao acessar a pasta principal: {erro}")

    return cmSubpastas # Retorna uma lista de objetos da classe Path

# Guarda o caminho dos resumos em uma lista
def pegarResumos(cmSubpastas):
    cmResumos = []

    for cmSubpasta in cmSubpastas:
        # Tenta acessar o caminho de uma subpasta
        try:
            itens = cmSubpasta.iterdir()

        # Em caso de qualquer erro, registra no log
        except Exception as erro:
            logging.error(f"Erro ao acessar {cmSubpasta}: {erro}")
            continue

        for item in itens:
            # Tenta acessar um caminho que está dentro de uma subpasta
            try:
                if item.is_file() and item.name.upper().startswith("RESUMO") and item.suffix == ".docx":
                    cmResumos.append(item)

            # Em caso de qualquer erro, registra no log
            except Exception as erro:
                logging.error(f"Erro ao acessar {item}: {erro}")
                continue

    return cmResumos       

# Abre os resumos e armazena seu conteúdo em uma lista
def lerResumos(cmResumos):
    docxResumos = []
    
    for cmResumo in cmResumos:
        
        # Tenta acessar o caminho de um resumo
        try:
            docxResumos.append(Document(cmResumo))

        # Em caso de qualquer erro, registra no log
        except Exception as erro:
            logging.error(f"Erro ao abrir {cmResumo.name}: {erro}")
            continue

    return docxResumos

# Cria um dicionário com as principais informações dos resumos
def extrairDados(docxResumos):
    dadosResumos = []

    for docxResumo in docxResumos:
        # Tenta abrir o resumo
        try:
            resumo = Document(docxResumo)

            # Define a estrutura do dicionário
            dados = {
                "Cliente": [],
                "Data_hora": [],
                "Orgao": [],
                "Objeto": [],
                "Modalidade": [],
                "Modo_disputa": [],
                "Critério_julgamento": [],
                "Fim_acolhimento": [],
                "Sistema": [],
                "Valor_referencial": []
            }

            # Percorre os parágrafos de um resumo
            for paragrafo in resumo.paragraphs:
                if "CLIENTE: " in paragrafo.text:
                    dados["Cliente"] = paragrafo.text.split("CLIENTE: ")[1].strip().upper()
                if "DATA E HORA: " in paragrafo.text:
                    dados["Data_hora"] = paragrafo.text.split("DATA E HORA: ")[1].strip()
                if "ÓRGÃO: " in paragrafo.text:
                    dados["Orgao"] = paragrafo.text.split("ÓRGÃO: ")[1].strip().upper()
                if "OBJETO: " in paragrafo.text:
                    dados["Objeto"] = paragrafo.text.split("OBJETO: ")[1].strip().upper()
                if "MODALIDADE: " in paragrafo.text:
                    dados["Modalidade"] = paragrafo.text.split("MODALIDADE: ")[1].strip().upper()
                if "MODO DE DISPUTA: " in paragrafo.text:
                    dados["Modo_disputa"] = paragrafo.text.split("MODO DE DISPUTA: ")[1].strip().upper()
                if "CRITÉRIO DE JULGAMENTO: " in paragrafo.text:
                    dados["Critério_julgamento"] = paragrafo.text.split("CRITÉRIO DE JULGAMENTO: ")[1].strip().upper()
                if "FIM DO ACOLHIMENTO DE PROPOSTA: " in paragrafo.text:
                    dados["Fim_acolhimento"] = paragrafo.text.split("FIM DO ACOLHIMENTO DE PROPOSTA: ")[1].strip().upper()
                if "SISTEMA: " in paragrafo.text:
                    dados["Sistema"] = paragrafo.text.split("SISTEMA: ")[1].strip().upper()
                if "VALOR REFERENCIAL: " in paragrafo.text:
                    dados["Valor_referencial"] = paragrafo.text.split("VALOR REFERENCIAL: ")[1].strip().upper()

        # Em caso de qualquer erro, registra no log
        except Exception as erro:
            logging.error(f"Erro ao ler arquivo: {erro}")
            continue

        return dados

######################################################################
# Main

# Cria um objeto (home) da classe Path para armazenar o diretório do usuário atual
home = Path.home() # Por exemplo C:\Users\fulano
cmPasta = home / "OneDrive" / "_APS LICITAÇÕES OPERACIONAL" / "00. DEMANDAS DA SEMANA"

# Armazena o caminho das subpastas de cmPasta em uma lista de Paths
cmSubpastas = pegarSubpastas(cmPasta)

# Verifica se a lista de subpastas está vazia
if not cmSubpastas:
    print("A pasta não contém subpastas.")

'''
for cmSubpasta in cmSubpastas:
    print(cmSubpasta.name)
'''

# Armazeno o caminho de cada resumo em uma lista de Paths
cmResumos = pegarResumos(cmSubpastas)

# Verifica se a lista de resumos está vazia
if not cmResumos:
    print("Não foram encontrados resumos.")

'''
for cmResumo in cmResumos:
    print(cmResumo.name)
'''

# Abre e lê todos os resumos encontrados
docxResumos = lerResumos(cmResumos)

'''
for paragrafo in docxResumos[0].paragraphs:
    print(paragrafo.text)
'''

dados = extrairDados(docxResumos)

for chave, valor in dados.items():
    print(f"{chave} -> {valor}\n")